from dataclasses import dataclass


@dataclass(frozen=True)
class ParsedDocument:
    text: str
    skipped_reason: str | None = None


def parse_document_bytes(content: bytes, mime_type: str | None, name: str) -> ParsedDocument:
    mime = mime_type or ""
    lower_name = name.lower()

    if mime.startswith("text/") or lower_name.endswith((".txt", ".md", ".csv")):
        return ParsedDocument(text=content.decode("utf-8", errors="ignore"))

    if mime == "application/pdf" or lower_name.endswith(".pdf"):
        return parse_pdf(content)

    if (
        mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        or lower_name.endswith(".docx")
    ):
        return parse_docx(content)

    return ParsedDocument(text="", skipped_reason=f"Unsupported document type: {mime or name}")


def parse_pdf(content: bytes) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ParsedDocument(text="", skipped_reason="Install pypdf to parse PDF files")

    import io

    reader = PdfReader(io.BytesIO(content))
    text = "\n\n".join(page.extract_text() or "" for page in reader.pages)
    return ParsedDocument(text=text)


def parse_docx(content: bytes) -> ParsedDocument:
    try:
        import docx
    except ImportError:
        return ParsedDocument(text="", skipped_reason="Install python-docx to parse DOCX files")

    import io

    document = docx.Document(io.BytesIO(content))
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    return ParsedDocument(text=text)
